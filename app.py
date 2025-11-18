#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
教师端AI备课助手Web应用（整合版本）
"""

from flask import Flask, render_template, request, jsonify, Response, send_file
from flask_cors import CORS
import os
import json
import argparse
from typing import List
from pathlib import Path
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 导入教师端备课模块
from teacher_planner import (
    collect_entities_llm,
    detect_intent_llm,
    call_hybrid_search,
    call_sports_meeting_search,
    generate_plan,
    generate_plan_stream,
    load_class_profiles,
)

# 导入班级数据分析模块
from analyze_class_data import (
    analyze_class_file,
    analyze_with_llm,
    analyze_uploaded_file,
    get_all_class_profiles,
    delete_class_profile,
    update_class_profile,
)

app = Flask(__name__)
CORS(app)

# 配置
SEARCH_BASE_URL = os.getenv("SEARCH_BASE_URL", "http://127.0.0.1:8001")


def gather_user_text(user_text: str, conversation_history: List[dict]) -> str:
    """收集用户输入和对话历史中的所有用户文本"""
    pieces: List[str] = []
    for msg in conversation_history or []:
        if msg.get("role") == "user":
            content = (msg.get("content") or "").strip()
            if content:
                pieces.append(content)
    if user_text:
        pieces.append(user_text.strip())
    return " ".join(pieces)


def detect_plan_type(current_text: str, conversation_history: List[dict]) -> str:
    """
    使用大模型进行意图识别，判断是全员运动会、课课练还是闲聊
    返回: "sports_meeting" | "lesson_plan" | "chat" | ""
    """
    try:
        intent = detect_intent_llm(current_text, conversation_history)
        # 如果是chat，返回空字符串（表示不是方案生成意图）
        if intent == "chat":
            return ""
        return intent
    except Exception as e:
        if os.getenv('DEBUG_AI','1')=='1':
            print(f"[TEACHER] 意图识别失败: {e}")
        return ""


def get_local_ip():
    """获取本机IP地址"""
    import socket
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except:
        return "127.0.0.1"


# ==================== 页面路由 ====================

@app.route('/')
def index():
    """主页：重定向到教师端"""
    return render_template('teacher.html')


@app.route('/teacher')
def teacher():
    """教师端AI备课助手页面"""
    return render_template('teacher.html')


@app.route('/class_data_manager')
def class_data_manager():
    """班级体测数据管理页面"""
    return render_template('class_data_manager.html')


# ==================== 教师端备课API ====================

@app.route('/api/teacher/plan', methods=['POST'])
def teacher_plan():
    """
    教师端AI备课助手：实体收集 -> 外部检索 -> 方案生成（非流式）
    
    入参：
      - message: 用户自然语言输入
      - conversation_history: 可选，对话历史记录
      - override_params: 可选，显式指定参数
    
    返回：
      - 如果缺少关键信息，返回 need_more_info=True 和 ask 提示语
      - 否则返回生成的方案
    """
    try:
        data = request.get_json() or {}
        user_text = (data.get('message') or '').strip()
        conversation_history = data.get('conversation_history') or []
        override_params = data.get('override_params') or {}
        
        if not user_text and not override_params:
            return jsonify({'success': False, 'message': '请提供message或override_params'}), 400

        if os.getenv('DEBUG_AI','1')=='1':
            print(f"[TEACHER] 收到请求: user_text={user_text}, history_len={len(conversation_history)}")

        # 使用大模型进行意图识别
        plan_type = detect_plan_type(user_text, conversation_history)
        is_sports_meeting = plan_type == "sports_meeting"
        is_lesson_plan = plan_type == "lesson_plan"
        is_chat = not plan_type
        
        if os.getenv('DEBUG_AI','1')=='1':
            print(f"[TEACHER] 意图识别: plan_type={plan_type}")

        # 实体抽取
        try:
            params, missing = collect_entities_llm(user_text, conversation_history)
            if os.getenv('DEBUG_AI','1')=='1':
                print(f"[TEACHER] 实体抽取: params={params}, missing={missing}")
        except Exception as e:
            if os.getenv('DEBUG_AI','1')=='1':
                print(f"[TEACHER] 实体抽取失败: {e}")
            return jsonify({'success': False, 'message': f'实体抽取失败: {e}'}), 500

        # 应用显式覆盖
        for k in ['semantic_query', 'count_query', 'grades_query', 'trained_weaknesses', 'top_k']:
            if k in override_params and override_params[k] not in (None, ''):
                params[k] = override_params[k]
                if k in missing:
                    missing.remove(k)

        # 添加意图类型到参数中
        params["plan_type"] = plan_type or ""
        params["conversation_history"] = conversation_history

        # 如果是闲聊，直接生成回复
        if is_chat:
            if os.getenv('DEBUG_AI','1')=='1':
                print("[TEACHER] 识别为闲聊，直接生成回复")
            response_text = generate_plan([], params, user_text, need_guidance=False)
            conversation_history.append({"role": "user", "content": user_text})
            conversation_history.append({"role": "assistant", "content": response_text})
            return jsonify({
                'success': True,
                'response': response_text,
                'conversation_history': conversation_history,
                'is_chat': True
            })

        # 判断是否需要引导
        need_guidance = False
        missing_fields = []

        if is_sports_meeting:
            # 全员运动会：需要操场条件、年级、人数等信息
            if not params.get("semantic_query") or not params.get("grades_query") or not params.get("count_query"):
                need_guidance = True
                if not params.get("semantic_query"):
                    missing_fields.append("semantic_query")
                if not params.get("grades_query"):
                    missing_fields.append("grades_query")
                if not params.get("count_query"):
                    missing_fields.append("count_query")
        elif is_lesson_plan:
            # 课课练：需要班级或薄弱项，满足任一即可
            has_grades = bool(params.get("grades_query"))
            has_weaknesses = bool(params.get("trained_weaknesses"))
            if not has_grades and not has_weaknesses:
                need_guidance = True
                missing_fields.extend(["grades_query", "trained_weaknesses"])

        # 如果需要引导，生成引导语
        if need_guidance:
            if os.getenv('DEBUG_AI','1')=='1':
                print(f"[TEACHER] 需要引导，缺失字段: {missing_fields}")
            guidance_text = generate_plan([], params, user_text, need_guidance=True)
            conversation_history.append({"role": "user", "content": user_text})
            conversation_history.append({"role": "assistant", "content": guidance_text})
            return jsonify({
                'success': True,
                'need_more_info': True,
                'ask': guidance_text,
                'conversation_history': conversation_history,
                'collected_params': params
            })

        # 调用检索接口
        try:
            if is_sports_meeting:
                # 全员运动会检索
                semantic_with_text = f"{params.get('semantic_query', '')} {user_text}".strip()
                payload = {
                    "semantic_query": semantic_with_text,
                    "count_query": str(params.get("count_query") or ""),
                    "grades_query": str(params.get("grades_query") or ""),
                    "top_k": int(params.get("top_k") or 5),
                }
                results = call_sports_meeting_search(SEARCH_BASE_URL, payload)
            else:
                # 课课练检索
                payload = {
                    "semantic_query": params.get("semantic_query") or "",
                    "count_query": str(params.get("count_query") or ""),
                    "grades_query": str(params.get("grades_query") or ""),
                    "trained_weaknesses": params.get("trained_weaknesses") or "",
                    "top_k": int(params.get("top_k") or 5),
                }
                results = call_hybrid_search(SEARCH_BASE_URL, payload)

            if os.getenv('DEBUG_AI','1')=='1':
                print(f"[TEACHER] 检索结果数量: {len(results)}")
                print("====== 检索结果原始数据 ======")
                print(json.dumps(results, ensure_ascii=False, indent=2))
                print("====== 检索结果结束 ======")
        except Exception as e:
            if os.getenv('DEBUG_AI','1')=='1':
                print(f"[TEACHER] 检索失败: {e}")
            results = []

        # 生成方案
        response_text = generate_plan(results, params, user_text, need_guidance=False)

        # 更新对话历史
        conversation_history.append({"role": "user", "content": user_text})
        conversation_history.append({"role": "assistant", "content": response_text})

        return jsonify({
            'success': True,
            'response': response_text,
            'conversation_history': conversation_history,
            'params': params,
            'results_count': len(results)
        })

    except Exception as e:
        if os.getenv('DEBUG_AI','1')=='1':
            print(f"[TEACHER] 教师端备课失败: {e}")
        return jsonify({'success': False, 'message': f'教师端备课失败: {str(e)}'}), 500


@app.route('/api/teacher/plan/stream', methods=['POST'])
def teacher_plan_stream():
    """
    教师端AI备课助手：流式输出方案

    入参：
      - message: 用户自然语言输入
      - conversation_history: 可选，对话历史记录
      - override_params: 可选，显式指定参数
    """
    try:
        data = request.get_json() or {}
        user_text = (data.get('message') or '').strip()
        conversation_history = data.get('conversation_history') or []
        override_params = data.get('override_params') or {}

        if not user_text and not override_params:
            return jsonify({'success': False, 'message': '请提供message或override_params'}), 400

        if os.getenv('DEBUG_AI','1')=='1':
            print(f"[TEACHER] 流式接口收到请求: user_text={user_text}")

        # 使用大模型进行意图识别
        plan_type = detect_plan_type(user_text, conversation_history)
        is_sports_meeting = plan_type == "sports_meeting"
        is_lesson_plan = plan_type == "lesson_plan"
        is_chat = not plan_type

        # 实体抽取
        params, missing = collect_entities_llm(user_text, conversation_history)

        # 应用显式覆盖
        for k in ['semantic_query', 'count_query', 'grades_query', 'trained_weaknesses', 'top_k']:
            if k in override_params and override_params[k] not in (None, ''):
                params[k] = override_params[k]

        # 添加意图类型到参数中
        params["plan_type"] = plan_type or ""
        params["conversation_history"] = conversation_history

        # 如果是闲聊，直接生成友好回复（流式）
        if is_chat:
            if os.getenv('DEBUG_AI','1')=='1':
                print("[TEACHER] 流式接口：识别为闲聊，直接生成回复")
            need_guidance = False
            missing_fields = []
        else:
            # 关键检查：根据场景判断是否需要更多信息
            count_query = params.get('count_query')
            grades_query = params.get('grades_query')
            semantic_query = params.get('semantic_query')
            trained_weaknesses_value = params.get('trained_weaknesses')

            missing_fields = []
            if is_sports_meeting:
                # 全员运动会：需要操场条件、年级、人数等信息
                if os.getenv('DEBUG_AI','1')=='1':
                    print(f"[TEACHER] 流式接口：全员运动会场景，检查必要字段 - semantic={semantic_query}, grades={grades_query}, count={count_query}")
                if not semantic_query:
                    missing_fields.append('semantic_query')
                if not grades_query:
                    missing_fields.append('grades_query')
                if not count_query:
                    missing_fields.append('count_query')
                if missing_fields and os.getenv('DEBUG_AI','1')=='1':
                    print("[TEACHER] 流式接口：⚠️ 全员运动会场景信息不全，进入引导流程，缺失=", missing_fields)
            elif is_lesson_plan:
                # 课课练：需要班级（grades_query）或弱项（trained_weaknesses），满足任一即可
                if os.getenv('DEBUG_AI','1')=='1':
                    print(f"[TEACHER] 流式接口：课课练场景，检查必要字段 - grades={grades_query}, trained_weaknesses={trained_weaknesses_value}")
                has_grades = bool(grades_query)
                has_weaknesses = bool(trained_weaknesses_value)
                if not has_grades and not has_weaknesses:
                    # 两个都没有，需要引导
                    missing_fields.append('grades_query_or_trained_weaknesses')
                if missing_fields and os.getenv('DEBUG_AI','1')=='1':
                    print("[TEACHER] 流式接口：⚠️ 课课练场景信息不全（缺少班级或弱项），进入引导流程")

            need_guidance = bool(missing_fields)

        # 如果是闲聊，直接生成友好回复（流式）
        if is_chat:
            def chat_stream():
                try:
                    model = OptimizedAIModel()
                    chat_messages = [
                        {"role": "system", "content": TEACHER_SYSTEM_PROMPT},
                        {"role": "user", "content": f"用户说：{user_text}\n\n请用友好、简洁的方式回复用户。如果是询问功能，可以介绍你可以帮助生成课课练备课方案和全员运动会方案。"}
                    ]
                    stream = model.client.chat.completions.create(
                        model=model.model,
                        messages=chat_messages,
                        max_tokens=200,
                        temperature=0.7,
                        stream=True,
                    )
                    for event in stream:
                        try:
                            delta = event.choices[0].delta
                            content = getattr(delta, "content", None)
                            if content:
                                yield content
                        except Exception:
                            chunk = None
                            try:
                                chunk = event["choices"][0]["delta"].get("content")
                            except Exception:
                                pass
                            if chunk:
                                yield chunk
                except Exception as e:
                    if os.getenv('DEBUG_AI','1')=='1':
                        print(f"[TEACHER] 流式接口：闲聊回复生成失败: {e}")
                    yield "您好！我是AI备课助理。我可以帮您生成课课练备课方案和全员运动会方案。请告诉我您的需求。"

            return Response(chat_stream(), mimetype='text/plain; charset=utf-8')

        if need_guidance:
            # 信息不全，使用generate_plan_stream生成引导语
            collected_so_far = {
                'semantic_query': params.get('semantic_query') or '',
                'count_query': params.get('count_query') or '',
                'grades_query': params.get('grades_query') or '',
                'trained_weaknesses': params.get('trained_weaknesses') or '',
                'plan_type': params.get('plan_type') or '',
                'top_k': int(params.get('top_k') or 5),
            }

            try:
                if os.getenv('DEBUG_AI','1')=='1':
                    print("[TEACHER] 流式接口：信息不全，调用generate_plan_stream生成引导语(流式)...")

                def guidance_stream():
                    ask_chunks = []
                    try:
                        for chunk in generate_plan_stream([], params, user_text, need_guidance=True):
                            ask_chunks.append(chunk)
                            yield chunk

                        final_ask = "".join(ask_chunks).strip()

                        if os.getenv('DEBUG_AI','1')=='1':
                            print("[TEACHER] 流式接口：引导语推送完成，长度=", len(final_ask))
                    except Exception as stream_err:
                        if os.getenv('DEBUG_AI','1')=='1':
                            print("[TEACHER] 流式接口：引导语流式推送异常:", stream_err)
                        yield f"[引导流错误] {stream_err}"

                resp = Response(guidance_stream(), mimetype='text/plain; charset=utf-8')
                resp.headers['X-Need-More-Info'] = '1'
                # HTTP响应头只能使用ASCII字符，需要将中文转义为\uXXXX格式
                resp.headers['X-Collected-Params'] = json.dumps(collected_so_far, ensure_ascii=True)
                return resp
            except Exception as e:
                if os.getenv('DEBUG_AI','1')=='1':
                    print(f"[TEACHER] 流式接口：引导语流式生成失败，使用兜底提示。错误: {e}")

                def fallback_stream():
                    yield "请说明需要重点提升的薄弱项（如：速度/力量/柔韧/灵敏/耐力/核心稳定/协调/平衡）"

                resp = Response(fallback_stream(), mimetype='text/plain; charset=utf-8')
                resp.headers['X-Need-More-Info'] = '1'
                # HTTP响应头只能使用ASCII字符，需要将中文转义为\uXXXX格式
                resp.headers['X-Collected-Params'] = json.dumps(collected_so_far, ensure_ascii=True)
                return resp

        # 调用检索接口
        try:
            if is_sports_meeting:
                # 全员运动会检索
                semantic_with_text = f"{params.get('semantic_query', '')} {user_text}".strip()
                payload = {
                    "semantic_query": semantic_with_text,
                    "count_query": str(params.get("count_query") or ""),
                    "grades_query": str(params.get("grades_query") or ""),
                    "top_k": int(params.get("top_k") or 5),
                }
                if os.getenv('DEBUG_AI','1')=='1':
                    print(f"[TEACHER] 流式接口：使用全员运动会检索 payload={payload}")
                results = call_sports_meeting_search(SEARCH_BASE_URL, payload)
                if os.getenv('DEBUG_AI','1')=='1':
                    print(f"[TEACHER] 流式接口：✅ 全员运动会检索成功，返回 {len(results)} 条")
                    print("====== 检索结果原始数据 ======")
                    print(json.dumps(results, ensure_ascii=False, indent=2))
                    print("====== 检索结果结束 ======")
            elif is_lesson_plan:
                # 课课练检索
                payload = {
                    "semantic_query": params.get("semantic_query") or "",
                    "count_query": str(params.get("count_query") or ""),
                    "grades_query": str(params.get("grades_query") or ""),
                    "trained_weaknesses": params.get("trained_weaknesses") or "",
                    "top_k": int(params.get("top_k") or 5),
                }
                if os.getenv('DEBUG_AI','1')=='1':
                    print(f"[TEACHER] 流式接口：调用检索 payload={payload}")
                    print(f"[TEACHER] 流式接口：🚀 开始调用检索接口 {SEARCH_BASE_URL}/extended-search/hybrid")
                results = call_hybrid_search(SEARCH_BASE_URL, payload)
                if os.getenv('DEBUG_AI','1')=='1':
                    print(f"[TEACHER] 流式接口：✅ 检索接口调用成功，返回 {len(results)} 条")
                    print("====== 检索结果原始数据 ======")
                    print(json.dumps(results, ensure_ascii=False, indent=2))
                    print("====== 检索结果结束 ======")
            else:
                results = []
        except Exception as e:
            if os.getenv('DEBUG_AI','1')=='1':
                print(f"[TEACHER] 流式接口：检索失败，使用空结果兜底: {e}")
            results = []

        # 流式生成方案
        def generate():
            try:
                for chunk in generate_plan_stream(results, params, user_text, need_guidance=False):
                    yield chunk
            except Exception as e:
                if os.getenv('DEBUG_AI','1')=='1':
                    print(f"[TEACHER] 流式生成失败: {e}")
                yield f"\n\n生成失败: {str(e)}"

        return Response(generate(), mimetype='text/plain; charset=utf-8')

    except Exception as e:
        if os.getenv('DEBUG_AI','1')=='1':
            print(f"[TEACHER] 流式生成失败: {e}")
        return jsonify({'success': False, 'message': f'流式生成失败: {str(e)}'}), 500


# ==================== 班级数据管理API ====================

@app.route('/api/class_data/upload', methods=['POST'])
def upload_class_data():
    """
    上传班级体测数据并分析（非流式）

    请求参数:
        - file: Excel文件
        - class_name: 班级名称

    返回:
        分析结果
    """
    try:
        # 检查是否有文件
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'message': '请上传文件'
            }), 400

        file = request.files['file']
        class_name = request.form.get('class_name', '').strip()

        if not class_name:
            return jsonify({
                'success': False,
                'message': '请提供班级名称'
            }), 400

        if file.filename == '':
            return jsonify({
                'success': False,
                'message': '请选择文件'
            }), 400

        # 读取文件内容
        file_content = file.read()

        # 分析数据
        result = analyze_uploaded_file(file_content, class_name)

        if result.get('success'):
            return jsonify({
                'success': True,
                'data': result
            })
        else:
            return jsonify({
                'success': False,
                'message': result.get('error', '分析失败')
            }), 500

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'分析失败: {str(e)}'
        }), 500


@app.route('/api/class_data/upload_stream', methods=['POST'])
def upload_class_data_stream():
    """
    上传班级体测数据并流式分析（使用大模型）

    请求参数:
        - file: Excel文件
        - class_name: 班级名称

    返回:
        流式分析过程
    """
    try:
        # 检查是否有文件
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'message': '请上传文件'
            }), 400

        file = request.files['file']
        class_name = request.form.get('class_name', '').strip()

        if not class_name:
            return jsonify({
                'success': False,
                'message': '请提供班级名称'
            }), 400

        if file.filename == '':
            return jsonify({
                'success': False,
                'message': '请选择文件'
            }), 400

        # 读取文件内容
        file_content = file.read()

        def generate():
            try:
                import pandas as pd
                import io

                # 读取Excel文件
                df = pd.read_excel(io.BytesIO(file_content))

                # 使用大模型流式分析
                profile = None
                for chunk in analyze_with_llm(df, class_name):
                    # 检查是否是最终的profile结果
                    if isinstance(chunk, tuple) and len(chunk) == 2 and chunk[0] == "__PROFILE__":
                        # 这是最终的profile字典
                        profile = chunk[1]
                    elif isinstance(chunk, str):
                        # 流式输出分析过程
                        yield f"data: {json.dumps({'type': 'progress', 'content': chunk}, ensure_ascii=False)}\n\n"

                # 保存配置
                if profile:
                    update_class_profile(class_name, profile)
                    yield f"data: {json.dumps({'type': 'success', 'profile': profile, 'message': '✅ 分析完成并已保存！'}, ensure_ascii=False)}\n\n"
                else:
                    yield f"data: {json.dumps({'type': 'error', 'message': '❌ 分析失败：未获取到分析结果'}, ensure_ascii=False)}\n\n"

                yield "data: [DONE]\n\n"

            except Exception as e:
                import traceback
                traceback.print_exc()
                yield f"data: {json.dumps({'type': 'error', 'message': f'❌ 分析失败: {str(e)}'}, ensure_ascii=False)}\n\n"
                yield "data: [DONE]\n\n"

        return Response(generate(), mimetype='text/event-stream')

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'分析失败: {str(e)}'
        }), 500


@app.route('/api/class_data/analyze/<path:filename>', methods=['POST'])
def analyze_existing_class_data(filename):
    """
    分析class_data文件夹中已有的体测数据文件

    参数:
        filename: 文件名（例如：一年级1班.xlsx）

    返回:
        分析结果
    """
    try:
        # 构建文件路径
        file_path = Path("class_data") / filename

        if not file_path.exists():
            return jsonify({
                'success': False,
                'message': f'文件不存在: {filename}'
            }), 404

        # 分析数据
        result = analyze_class_file(file_path)

        return jsonify({
            'success': True,
            'data': result
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'分析失败: {str(e)}'
        }), 500


@app.route('/api/class_data/profiles', methods=['GET'])
def get_class_profiles_api():
    """
    获取所有班级配置

    返回:
        所有班级配置
    """
    try:
        profiles = get_all_class_profiles()
        return jsonify({
            "success": True,
            "data": profiles,
            "count": len(profiles)
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"获取班级配置失败: {str(e)}"
        }), 500


@app.route('/api/class_data/profile/<class_name>', methods=['DELETE'])
def delete_class_profile_api(class_name):
    """
    删除班级配置

    参数:
        class_name: 班级名称

    返回:
        删除结果
    """
    try:
        success = delete_class_profile(class_name)

        if success:
            return jsonify({
                "success": True,
                "message": f"已删除班级配置: {class_name}"
            })
        else:
            return jsonify({
                "success": False,
                "message": f"班级配置不存在: {class_name}"
            }), 404

    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"删除失败: {str(e)}"
        }), 500


@app.route('/api/class_data/download/<class_name>', methods=['GET'])
def download_class_excel(class_name):
    """
    下载班级配置的Excel文件

    参数:
        class_name: 班级名称

    返回:
        Excel文件
    """
    try:
        # 获取班级配置
        profiles = get_all_class_profiles()

        if class_name not in profiles:
            return jsonify({
                "success": False,
                "message": f"班级配置不存在: {class_name}"
            }), 404

        profile = profiles[class_name]

        # 创建Excel文件
        output = BytesIO()

        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Sheet 1: 班级基本信息
            basic_info = {
                '班级名称': [class_name],
                '年级': [f"{profile.get('grades_query', '')}年级"],
                '薄弱项': [', '.join(profile.get('weaknesses', []))],
                '描述': [profile.get('description', '')]
            }
            df_basic = pd.DataFrame(basic_info)
            df_basic.to_excel(writer, sheet_name='班级信息', index=False)

            # Sheet 2: 学生分组详情
            if 'student_groups' in profile and profile['student_groups']:
                all_students = []

                for group_key, group_info in profile['student_groups'].items():
                    weakness_items = ', '.join(group_info.get('weakness_items', []))

                    if 'student_details' in group_info and group_info['student_details']:
                        for student in group_info['student_details']:
                            # 优先获取学生编号，尝试多个可能的字段
                            student_id = student.get('学生编号', '') or student.get('学号', '') or student.get('编号', '')
                            student_name = student.get('姓名', '')
                            student_index = student.get('序号', '')

                            # 如果没有姓名，使用"学生X"
                            if not student_name and student_index:
                                student_name = f'学生{student_index}'

                            student_row = {
                                '分组': group_key,
                                '薄弱项目': weakness_items,
                                '学生编号': str(student_id) if student_id else '',
                                '姓名': student_name,
                                '性别': student.get('性别', '')
                            }
                            all_students.append(student_row)

                if all_students:
                    df_students = pd.DataFrame(all_students)
                    df_students.to_excel(writer, sheet_name='学生分组', index=False)

            # Sheet 3: 各项体测统计（如果有的话）
            if 'test_stats' in profile and profile['test_stats']:
                stats_data = []
                for item, stats in profile['test_stats'].items():
                    stats_row = {
                        '体测项目': item,
                        '维度': stats.get('dimension', ''),
                        '优秀人数': stats.get('excellent', 0),
                        '良好人数': stats.get('good', 0),
                        '及格人数': stats.get('pass', 0),
                        '不及格人数': stats.get('fail', 0)
                    }
                    stats_data.append(stats_row)

                if stats_data:
                    df_stats = pd.DataFrame(stats_data)
                    df_stats.to_excel(writer, sheet_name='体测统计', index=False)

        output.seek(0)

        # 返回Excel文件
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'{class_name}_配置.xlsx'
        )

    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"下载失败: {str(e)}"
        }), 500


@app.route('/api/class_data/download_word/<class_name>', methods=['GET'])
def download_class_word(class_name):
    """
    下载班级配置的Word文档（美化版）

    参数:
        class_name: 班级名称

    返回:
        Word文件
    """
    try:
        # 获取班级配置
        profiles = get_all_class_profiles()

        if class_name not in profiles:
            return jsonify({
                "success": False,
                "message": f"班级配置不存在: {class_name}"
            }), 404

        profile = profiles[class_name]

        # 创建Word文档
        doc = Document()

        # 设置默认字体为中文字体（解决乱码问题）
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 添加标题
        title = doc.add_heading(f'{class_name} 体测数据分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title.runs:
            title.runs[0].font.size = Pt(20)
            title.runs[0].font.name = '宋体'
            title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 添加空行
        doc.add_paragraph()

        # 1. 班级基本信息
        heading1 = doc.add_heading('一、班级基本信息', 1)
        for run in heading1.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 创建基本信息表格
        table1 = doc.add_table(rows=4, cols=2)
        table1.style = 'Light Grid Accent 1'

        # 设置表头
        cells = table1.rows[0].cells
        cells[0].text = '班级名称'
        cells[1].text = class_name

        cells = table1.rows[1].cells
        cells[0].text = '年级'
        cells[1].text = f"{profile.get('grades_query', '')}年级"

        cells = table1.rows[2].cells
        cells[0].text = '薄弱项'
        cells[1].text = ', '.join(profile.get('weaknesses', []))

        cells = table1.rows[3].cells
        cells[0].text = '分析描述'
        cells[1].text = profile.get('description', '')

        # 设置表格样式（添加字体设置）
        for row in table1.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 第一列加粗
            if row.cells[0].paragraphs and row.cells[0].paragraphs[0].runs:
                row.cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

        # 2. 学生分组详情
        if 'student_groups' in profile and profile['student_groups']:
            heading2 = doc.add_heading('二、学生分组详情', 1)
            for run in heading2.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            for group_key, group_info in profile['student_groups'].items():
                # 分组标题
                group_heading = doc.add_heading(f'{group_key}薄弱组（{group_info["count"]}人）', 2)
                for run in group_heading.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                # 薄弱项目
                weakness_items = ', '.join(group_info.get('weakness_items', []))
                p = doc.add_paragraph(f'体测不及格项目：{weakness_items}')
                for run in p.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                # 学生列表表格
                if 'student_details' in group_info and group_info['student_details']:
                    students = group_info['student_details']

                    # 创建表格（表头 + 学生行）
                    table2 = doc.add_table(rows=len(students) + 1, cols=3)
                    table2.style = 'Light List Accent 1'

                    # 表头
                    header_cells = table2.rows[0].cells
                    header_cells[0].text = '学生编号'
                    header_cells[1].text = '姓名'
                    header_cells[2].text = '性别'

                    # 表头样式
                    for cell in header_cells:
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            cell.paragraphs[0].runs[0].font.bold = True
                            cell.paragraphs[0].runs[0].font.size = Pt(11)
                            cell.paragraphs[0].runs[0].font.name = '宋体'
                            cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    # 填充学生数据
                    for i, student in enumerate(students, start=1):
                        row_cells = table2.rows[i].cells
                        # 优先获取学生编号，尝试多个可能的字段
                        student_id = student.get('学生编号', '') or student.get('学号', '') or student.get('编号', '')
                        student_name = student.get('姓名', '')
                        student_index = student.get('序号', '')

                        # 如果没有姓名，使用"学生X"
                        if not student_name and student_index:
                            student_name = f'学生{student_index}'

                        row_cells[0].text = str(student_id) if student_id else ''
                        row_cells[1].text = student_name
                        row_cells[2].text = student.get('性别', '')

                        # 设置字体
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = '宋体'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                doc.add_paragraph()

        # 3. 体测统计
        if 'test_stats' in profile and profile['test_stats']:
            heading3 = doc.add_heading('三、体测统计', 1)
            for run in heading3.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            stats = profile['test_stats']

            # 创建统计表格
            table3 = doc.add_table(rows=len(stats) + 1, cols=6)
            table3.style = 'Light Grid Accent 1'

            # 表头
            header_cells = table3.rows[0].cells
            headers = ['体测项目', '维度', '优秀人数', '良好人数', '及格人数', '不及格人数']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                if header_cells[i].paragraphs and header_cells[i].paragraphs[0].runs:
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
                    header_cells[i].paragraphs[0].runs[0].font.name = '宋体'
                    header_cells[i].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 填充数据
            for i, (item, stat) in enumerate(stats.items(), start=1):
                row_cells = table3.rows[i].cells
                row_cells[0].text = item
                row_cells[1].text = stat.get('dimension', '')
                row_cells[2].text = str(stat.get('excellent', 0))
                row_cells[3].text = str(stat.get('good', 0))
                row_cells[4].text = str(stat.get('pass', 0))
                row_cells[5].text = str(stat.get('fail', 0))

                # 设置字体
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 保存到BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        # 返回Word文件
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{class_name}_配置.docx'
        )

    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"下载失败: {str(e)}"
        }), 500


@app.route('/api/class_data/batch_analyze', methods=['POST'])
def batch_analyze_class_data():
    """
    批量分析class_data文件夹中的所有体测数据

    请求参数:
        - max_count: 最多分析多少个班级（可选）

    返回:
        批量分析结果
    """
    try:
        data = request.get_json() or {}
        max_count = data.get('max_count', None)

        # 获取class_data文件夹中的所有Excel文件
        class_data_dir = Path("class_data")
        if not class_data_dir.exists():
            return jsonify({
                'success': False,
                'message': 'class_data文件夹不存在'
            }), 404

        excel_files = list(class_data_dir.glob("*.xlsx")) + list(class_data_dir.glob("*.xls"))

        if max_count:
            excel_files = excel_files[:max_count]

        results = []
        for file_path in excel_files:
            try:
                result = analyze_class_file(file_path)
                results.append({
                    'class_name': file_path.stem,
                    'success': True,
                    'data': result
                })
            except Exception as e:
                results.append({
                    'class_name': file_path.stem,
                    'success': False,
                    'error': str(e)
                })

        return jsonify({
            'success': True,
            'total': len(excel_files),
            'results': results
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'批量分析失败: {str(e)}'
        }), 500


# ==================== 主程序入口 ====================

if __name__ == '__main__':
    # 解析命令行参数
    parser = argparse.ArgumentParser(description='教师端AI备课助手Web应用')
    parser.add_argument('--host', '-H', type=str, default=None,
                        help='监听地址 (默认: 0.0.0.0，允许所有IP访问)')
    parser.add_argument('--port', '-p', type=int, default=None,
                        help='端口号 (默认: 5000)')
    parser.add_argument('--debug', action='store_true',
                        help='开启调试模式')
    parser.add_argument('--no-debug', dest='debug', action='store_false',
                        help='关闭调试模式')
    parser.set_defaults(debug=None)  # 默认不设置，使用环境变量或默认值
    
    args = parser.parse_args()
    
    # 优先级：命令行参数 > 环境变量 > 默认值
    host = args.host if args.host is not None else os.getenv('HOST', '0.0.0.0')
    port = args.port if args.port is not None else int(os.getenv('PORT', 5000))
    
    # 调试模式：命令行参数 > 环境变量 > 默认值
    if args.debug is not None:
        debug = args.debug
    else:
        debug_env = os.getenv('DEBUG', 'True')
        debug = debug_env.lower() == 'true'

    local_ip = get_local_ip()

    print("=" * 60)
    print("教师端AI备课助手（整合版本）")
    print("=" * 60)
    print(f"监听地址: {host}:{port}")
    print(f"本地访问: http://127.0.0.1:{port}")
    print(f"局域网访问: http://{local_ip}:{port}")
    print(f"调试模式: {'开启' if debug else '关闭'}")
    print("=" * 60)
    print("\n可用页面:")
    print(f"  - 教师端备课助手: http://127.0.0.1:{port}/teacher")
    print(f"  - 班级数据管理: http://127.0.0.1:{port}/class_data_manager")
    print("=" * 60)

    app.run(
        host=host,
        port=port,
        debug=debug
    )

